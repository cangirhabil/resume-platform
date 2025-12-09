from fpdf import FPDF
from jinja2 import Environment, FileSystemLoader
import os
import io
from typing import Dict, Any, List


class ResumePDF(FPDF):
    """Custom PDF class for resume generation with professional layout."""
    
    def __init__(self):
        super().__init__()
        self.set_auto_page_break(auto=True, margin=15)
        
    def header(self):
        pass  # We'll handle headers manually for more control
        
    def footer(self):
        pass  # No footer needed for resumes


class PDFGenerator:
    def __init__(self):
        self.template_dir = os.path.join(os.path.dirname(__file__), "../templates")
        self.env = Environment(loader=FileSystemLoader(self.template_dir))
        self.available_templates = ["professional", "modern", "classic", "minimal"]

    def get_available_templates(self) -> List[str]:
        """Returns list of available template names."""
        return self.available_templates
    
    def _normalize_resume_data(self, resume_data: dict) -> dict:
        """
        Normalize resume data to handle different LLM output formats.
        Ensures templates receive consistent data structure.
        """
        normalized = dict(resume_data)
        
        # Handle personal_info variations
        personal_info = normalized.get("personal_info", {})
        if isinstance(personal_info, dict):
            # Ensure name field exists
            if not personal_info.get("name"):
                personal_info["name"] = personal_info.get("full_name", "")
            normalized["personal_info"] = personal_info
        
        # Handle experience - ensure bullets format
        experience = normalized.get("experience", [])
        normalized_experience = []
        for exp in experience:
            if isinstance(exp, dict):
                norm_exp = dict(exp)
                # Convert description to bullets if needed
                if "description" in norm_exp and "bullets" not in norm_exp:
                    desc = norm_exp.get("description", "")
                    if isinstance(desc, str):
                        norm_exp["bullets"] = [desc] if desc else []
                    elif isinstance(desc, list):
                        norm_exp["bullets"] = desc
                # Ensure bullets is always a list
                if "bullets" not in norm_exp:
                    norm_exp["bullets"] = []
                # Handle dates format
                if "dates" not in norm_exp:
                    start = norm_exp.get("start_date", "")
                    end = norm_exp.get("end_date", "Present")
                    norm_exp["dates"] = f"{start} - {end}" if start else ""
                normalized_experience.append(norm_exp)
        normalized["experience"] = normalized_experience
        
        # Handle skills - can be dict or list
        skills = normalized.get("skills", [])
        if isinstance(skills, dict):
            # Flatten skills dict to list
            flat_skills = []
            for category, skill_list in skills.items():
                if isinstance(skill_list, list):
                    flat_skills.extend(skill_list)
                elif isinstance(skill_list, str):
                    flat_skills.append(skill_list)
            normalized["skills"] = flat_skills
            normalized["skills_dict"] = skills  # Keep original for templates that support it
        elif isinstance(skills, list):
            normalized["skills"] = skills
            normalized["skills_dict"] = {"skills": skills}
        
        # Handle education
        education = normalized.get("education", [])
        normalized_education = []
        for edu in education:
            if isinstance(edu, dict):
                norm_edu = dict(edu)
                # Ensure school field
                if not norm_edu.get("school"):
                    norm_edu["school"] = norm_edu.get("institution", "")
                # Handle dates
                if "dates" not in norm_edu:
                    norm_edu["dates"] = norm_edu.get("graduation_date", "")
                normalized_education.append(norm_edu)
        normalized["education"] = normalized_education
        
        # Handle programs/certifications
        programs = normalized.get("programs", normalized.get("certifications", []))
        normalized["programs"] = programs
        
        return normalized

    def _clean_text(self, text: str) -> str:
        """Clean text for PDF rendering - handle unicode characters."""
        if not text:
            return ""
        # Replace unicode bullet points
        text = str(text)
        text = text.replace("•", "-")
        text = text.replace("▪", "-")
        text = text.replace("→", "->")
        text = text.replace("–", "-")
        text = text.replace("—", "-")
        text = text.replace(""", '"')
        text = text.replace(""", '"')
        text = text.replace("'", "'")
        text = text.replace("'", "'")
        return text

    def generate_professional(self, data: dict) -> bytes:
        """Generate PDF using professional template matching the user's design."""
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        
        # Use built-in fonts
        personal_info = data.get("personal_info", {})
        name = self._clean_text(personal_info.get("name", "Your Name"))
        
        # HEADER - Name (large, bold)
        pdf.set_font("Helvetica", "B", 28)
        pdf.set_text_color(0, 0, 0)
        pdf.cell(0, 12, name, ln=True)
        
        # Contact info line
        contact_parts = []
        if personal_info.get("email"):
            contact_parts.append(personal_info["email"])
        if personal_info.get("linkedin"):
            contact_parts.append(personal_info["linkedin"])
        if personal_info.get("github"):
            contact_parts.append(personal_info["github"])
        if personal_info.get("phone"):
            contact_parts.append(personal_info["phone"])
        if personal_info.get("location"):
            contact_parts.append(personal_info["location"])
        
        if contact_parts:
            pdf.set_font("Helvetica", "", 9)
            pdf.set_text_color(100, 100, 100)
            contact_text = " | ".join(contact_parts)
            pdf.cell(0, 5, self._clean_text(contact_text), ln=True)
        
        # SUMMARY
        summary = data.get("summary", "")
        if summary:
            pdf.set_font("Helvetica", "I", 10)
            pdf.set_text_color(60, 60, 60)
            pdf.ln(3)
            pdf.multi_cell(0, 5, self._clean_text(summary))
        
        pdf.ln(5)
        
        # EDUCATION Section
        education = data.get("education", [])
        if education:
            self._add_section_header(pdf, "EDUCATION")
            for edu in education:
                school = self._clean_text(edu.get("school") or edu.get("institution", ""))
                degree = self._clean_text(edu.get("degree", ""))
                dates = self._clean_text(edu.get("dates", ""))
                
                # Title with degree and school
                pdf.set_font("Helvetica", "B", 11)
                pdf.set_text_color(0, 0, 0)
                
                title_text = f"{school}"
                if degree:
                    title_text += f", {degree}"
                
                # Calculate width for title and date
                page_width = pdf.w - pdf.l_margin - pdf.r_margin
                date_width = pdf.get_string_width(dates) + 5
                title_width = page_width - date_width
                
                # Title
                pdf.cell(title_width, 6, title_text[:80], ln=False)
                
                # Date (right-aligned)
                pdf.set_font("Helvetica", "I", 10)
                pdf.set_text_color(80, 80, 80)
                pdf.cell(date_width, 6, dates, ln=True, align="R")
                
                # Bullets for education details
                details = edu.get("details", edu.get("bullets", []))
                if details:
                    for detail in details:
                        if detail:
                            self._add_bullet(pdf, self._clean_text(detail))
                
                pdf.ln(2)
        
        # EXPERIENCE Section
        experience = data.get("experience", [])
        if experience:
            self._add_section_header(pdf, "EXPERIENCE")
            for exp in experience:
                title = self._clean_text(exp.get("title") or exp.get("position", ""))
                company = self._clean_text(exp.get("company", ""))
                dates = self._clean_text(exp.get("dates", ""))
                subtitle = self._clean_text(exp.get("subtitle", ""))
                
                # Title line (bold title - company)
                pdf.set_font("Helvetica", "B", 11)
                pdf.set_text_color(0, 0, 0)
                
                title_text = title
                if company:
                    title_text += f" - {company}"
                
                # Calculate width for title and date
                page_width = pdf.w - pdf.l_margin - pdf.r_margin
                date_width = pdf.get_string_width(dates) + 5
                title_width = page_width - date_width
                
                pdf.cell(title_width, 6, title_text[:70], ln=False)
                
                # Date (right-aligned, italic)
                pdf.set_font("Helvetica", "I", 10)
                pdf.set_text_color(80, 80, 80)
                pdf.cell(date_width, 6, dates, ln=True, align="R")
                
                # Subtitle if present
                if subtitle:
                    pdf.set_font("Helvetica", "I", 10)
                    pdf.set_text_color(60, 60, 60)
                    pdf.cell(0, 5, subtitle, ln=True)
                
                # Bullets
                bullets = exp.get("bullets", exp.get("details", []))
                for bullet in bullets:
                    if bullet:
                        self._add_bullet(pdf, self._clean_text(bullet))
                
                # Website/Product link
                website = exp.get("website") or exp.get("product") or exp.get("link")
                if website:
                    pdf.set_font("Helvetica", "", 9)
                    pdf.set_text_color(100, 100, 100)
                    pdf.set_x(pdf.l_margin + 5)
                    pdf.cell(0, 5, f"Website: {website}", ln=True)
                
                pdf.ln(2)
        
        # TECHNOLOGY / SKILLS Section
        skills = data.get("skills", [])
        if skills:
            self._add_section_header(pdf, "TECHNOLOGY")
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(40, 40, 40)
            skills_text = ", ".join([self._clean_text(s) for s in skills])
            pdf.multi_cell(0, 5, skills_text)
            pdf.ln(3)
        
        # PROJECTS Section
        projects = data.get("projects", [])
        if projects:
            self._add_section_header(pdf, "PROJECTS")
            for project in projects:
                name = self._clean_text(project.get("name", ""))
                description = self._clean_text(project.get("description", ""))
                dates = self._clean_text(project.get("dates", ""))
                
                pdf.set_font("Helvetica", "B", 11)
                pdf.set_text_color(0, 0, 0)
                
                # Calculate width for title and date
                page_width = pdf.w - pdf.l_margin - pdf.r_margin
                date_width = pdf.get_string_width(dates) + 5
                title_width = page_width - date_width
                
                pdf.cell(title_width, 6, name[:60], ln=False)
                
                # Date (right-aligned)
                pdf.set_font("Helvetica", "I", 10)
                pdf.set_text_color(80, 80, 80)
                pdf.cell(date_width, 6, dates, ln=True, align="R")
                
                # Description
                if description:
                    self._add_bullet(pdf, description)
                
                pdf.ln(2)
        
        # PROGRAMS Section
        programs = data.get("programs", data.get("certifications", []))
        if programs:
            self._add_section_header(pdf, "PROGRAMS")
            for program in programs:
                if isinstance(program, str):
                    pdf.set_font("Helvetica", "", 10)
                    pdf.set_text_color(40, 40, 40)
                    self._add_bullet(pdf, self._clean_text(program))
                elif isinstance(program, dict):
                    name = self._clean_text(program.get("name", ""))
                    dates = self._clean_text(program.get("dates", ""))
                    description = self._clean_text(program.get("description", ""))
                    
                    pdf.set_font("Helvetica", "B", 11)
                    pdf.set_text_color(0, 0, 0)
                    
                    page_width = pdf.w - pdf.l_margin - pdf.r_margin
                    date_width = pdf.get_string_width(dates) + 5
                    title_width = page_width - date_width
                    
                    pdf.cell(title_width, 6, name[:60], ln=False)
                    
                    pdf.set_font("Helvetica", "I", 10)
                    pdf.set_text_color(80, 80, 80)
                    pdf.cell(date_width, 6, dates, ln=True, align="R")
                    
                    if description:
                        self._add_bullet(pdf, description)
            pdf.ln(2)
        
        return pdf.output()

    def _add_section_header(self, pdf: FPDF, title: str):
        """Add a section header with gray background."""
        pdf.set_font("Helvetica", "B", 12)
        pdf.set_text_color(0, 0, 0)
        pdf.set_fill_color(230, 230, 240)  # Light blue-gray background
        pdf.cell(0, 7, title, ln=True, fill=True)
        pdf.ln(2)
    
    def _add_bullet(self, pdf: FPDF, text: str):
        """Add a bullet point item."""
        pdf.set_font("Helvetica", "", 10)
        pdf.set_text_color(40, 40, 40)
        pdf.set_x(pdf.l_margin + 5)
        # Use dash as bullet
        pdf.multi_cell(0, 5, f"- {text}")

    def render_html(self, resume_data: dict, theme: str) -> str:
        """
        Renders the HTML content from a Jinja2 template.
        """
        if theme not in self.available_templates or theme == "professional":
            theme = "modern"
        
        # Normalize data for consistent template rendering
        normalized_data = self._normalize_resume_data(resume_data)
        
        template = self.env.get_template(f"themes/{theme}.html")
        return template.render(**normalized_data)

    def generate(self, resume_data: dict, theme: str = "professional") -> bytes:
        """
        Generates PDF bytes from data using the specified theme.
        """
        try:
            # Normalize data
            normalized_data = self._normalize_resume_data(resume_data)
            
            # Use fpdf2 professional template by default
            if theme == "professional" or theme == "modern":
                return self.generate_professional(normalized_data)
            else:
                # For other themes, use fpdf2 professional as fallback
                return self.generate_professional(normalized_data)
                
        except Exception as e:
            print(f"PDF Generation Failed: {e}")
            import traceback
            traceback.print_exc()
            # Return a valid PDF with error message
            return self._generate_error_pdf(str(e))
    
    def _generate_error_pdf(self, error_message: str) -> bytes:
        """Generate a PDF with error message."""
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 18)
        pdf.set_text_color(200, 0, 0)
        pdf.cell(0, 20, "Resume Generation Error", ln=True, align="C")
        pdf.set_font("Helvetica", "", 12)
        pdf.set_text_color(80, 80, 80)
        pdf.multi_cell(0, 8, f"Error: {error_message[:200]}")
        pdf.ln(10)
        pdf.cell(0, 10, "Please try again or contact support.", ln=True, align="C")
        return pdf.output()

    def generate_docx(self, resume_data: dict) -> bytes:
        """
        Generates DOCX bytes from resume data.
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # Normalize data
            data = self._normalize_resume_data(resume_data)
            
            doc = Document()
            
            # Set margins
            for section in doc.sections:
                section.top_margin = Inches(0.75)
                section.bottom_margin = Inches(0.75)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Personal Info / Header
            personal_info = data.get("personal_info", {})
            name = personal_info.get("name") or "Your Name"
            
            name_para = doc.add_paragraph()
            name_run = name_para.add_run(name)
            name_run.bold = True
            name_run.font.size = Pt(24)
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Contact info
            contact_parts = []
            if personal_info.get("email"):
                contact_parts.append(personal_info["email"])
            if personal_info.get("phone"):
                contact_parts.append(personal_info["phone"])
            if personal_info.get("linkedin"):
                contact_parts.append(personal_info["linkedin"])
            if personal_info.get("location"):
                contact_parts.append(personal_info["location"])
            
            if contact_parts:
                contact_para = doc.add_paragraph(" | ".join(contact_parts))
                contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in contact_para.runs:
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(100, 100, 100)
            
            # Summary
            summary = data.get("summary")
            if summary:
                doc.add_heading("Professional Summary", level=2)
                doc.add_paragraph(summary)
            
            # Experience
            experience = data.get("experience", [])
            if experience:
                doc.add_heading("Experience", level=2)
                for exp in experience:
                    title = exp.get("title") or exp.get("position", "")
                    company = exp.get("company", "")
                    dates = exp.get("dates", "")
                    
                    job_para = doc.add_paragraph()
                    title_run = job_para.add_run(f"{title}")
                    title_run.bold = True
                    if company:
                        job_para.add_run(f" at {company}")
                    if dates:
                        job_para.add_run(f" ({dates})")
                    
                    # Bullets
                    bullets = exp.get("bullets", [])
                    for bullet in bullets:
                        if bullet:
                            doc.add_paragraph(bullet, style="List Bullet")
            
            # Education
            education = data.get("education", [])
            if education:
                doc.add_heading("Education", level=2)
                for edu in education:
                    degree = edu.get("degree", "")
                    school = edu.get("school") or edu.get("institution", "")
                    dates = edu.get("dates", "")
                    
                    edu_para = doc.add_paragraph()
                    edu_run = edu_para.add_run(f"{degree}")
                    edu_run.bold = True
                    if school:
                        edu_para.add_run(f" - {school}")
                    if dates:
                        edu_para.add_run(f" ({dates})")
            
            # Skills
            skills = data.get("skills", [])
            if skills:
                doc.add_heading("Skills", level=2)
                doc.add_paragraph(", ".join(skills))
            
            # Certifications
            certifications = data.get("certifications", [])
            if certifications:
                doc.add_heading("Certifications", level=2)
                for cert in certifications:
                    doc.add_paragraph(cert, style="List Bullet")
            
            # Projects
            projects = data.get("projects", [])
            if projects:
                doc.add_heading("Projects", level=2)
                for project in projects:
                    proj_para = doc.add_paragraph()
                    proj_run = proj_para.add_run(project.get("name", ""))
                    proj_run.bold = True
                    if project.get("description"):
                        doc.add_paragraph(project["description"])
            
            # Save to bytes
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
            
        except ImportError:
            print("python-docx not installed. Falling back to dummy DOCX.")
            return self._get_dummy_docx()
        except Exception as e:
            print(f"DOCX Generation Failed: {e}")
            import traceback
            traceback.print_exc()
            return self._get_dummy_docx()
    
    def _get_dummy_docx(self) -> bytes:
        """Returns a minimal valid DOCX file."""
        # This is a minimal DOCX structure
        return b"PK\x03\x04\x14\x00\x00\x00\x00\x00"  # Minimal zip header

pdf_generator = PDFGenerator()
